from __future__ import annotations

import io
from datetime import datetime
from typing import Any, Dict, List

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


def generate_docx_report(
    differences: List[Dict[str, Any]],
    old_version: str = "v1.0",
    new_version: str = "v1.1",
    reviewer_name: str = "維修工程與安全品質審查員",
) -> bytes:
    """Generate a formal DOCX version difference review report."""
    doc = Document()

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("AI 船舶技術文件版本差異審查報告")
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(14, 118, 110)  # #0e766e

    # Meta Info
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.add_run(f"舊版：{old_version}   |   新版：{new_version}   |   報告時間：{datetime.now().strftime('%Y-%m-%d %H:%M')}\n")

    # Safety Warning Banner
    unreviewed_high = [d for d in differences if d.get("risk") == "High" and d.get("review_status") == "未覆核"]
    if unreviewed_high:
        warn_p = doc.add_paragraph()
        warn_run = warn_p.add_run(f"⚠️ 警告：尚有 {len(unreviewed_high)} 項 High 重大安全變更尚未完成人工覆核！本報告為修訂草稿。")
        warn_run.font.bold = True
        warn_run.font.color.rgb = RGBColor(180, 59, 55)

    # Executive Summary Card
    doc.add_heading("一、 變更摘要統計 (Executive Summary)", level=1)
    high_count = sum(d.get("risk") == "High" for d in differences)
    reviewed_count = sum(d.get("review_status") != "未覆核" for d in differences)

    summary_p = doc.add_paragraph()
    summary_p.add_run(f"• 總差異筆數：{len(differences)} 筆\n")
    summary_p.add_run(f"• 重大安全變更 (High Risk)：{high_count} 筆\n")
    summary_p.add_run(f"• 人工覆核達成率：{reviewed_count} / {len(differences)} ({(reviewed_count/len(differences)*100 if differences else 0):.1f}%)\n")

    # Audit Trail Table
    doc.add_heading("二、 可追溯差異明細與審查對照 (Traceable Audit Trail)", level=1)

    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    headers = ["ID", "類型", "風險", "變更解讀與可能影響", "舊版來源", "新版來源", "人工簽核與筆記"]

    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    for item in differences:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.get("id", "-"))
        row_cells[1].text = str(item.get("change_type", "-"))
        row_cells[2].text = str(item.get("risk", "Low"))
        row_cells[3].text = f"{item.get('explanation', '')}\n[受影響]: {item.get('affected', '')}"

        old_src = item.get("old") or {}
        new_src = item.get("new") or {}

        row_cells[4].text = f"p.{old_src.get('page', '-')}\n{old_src.get('text', '-')[:60]}..."
        row_cells[5].text = f"p.{new_src.get('page', '-')}\n{new_src.get('text', '-')[:60]}..."
        row_cells[6].text = f"狀態: {item.get('review_status', '未覆核')}\n筆記: {item.get('reviewer_note', '-')}"

    # Sign-off Section
    doc.add_heading("三、 人工審查簽核欄 (Human Sign-off Block)", level=1)
    sign_p = doc.add_paragraph()
    sign_p.add_run(f"審查人員： __________________ (簽名)\n\n")
    sign_p.add_run(f"安品與技術主管： __________________ (簽名)\n\n")
    sign_p.add_run(f"簽核日期： 年   月   日\n")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
